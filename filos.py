import streamlit as st
import requests
from docx import Document
from docx.shared import Pt
from io import BytesIO
import time
from bs4 import BeautifulSoup
import re
import PyPDF2

# Configuración de la página
st.set_page_config(
    page_title="Adaptador de Obras Filosóficas",
    layout="wide",
    initial_sidebar_state="expanded",
)

# Título de la aplicación
st.title("Adaptador de Obras Filosóficas para Estudiantes de 16 Años")

# Descripción
st.markdown("""
Esta aplicación adapta una obra filosófica proporcionada en **PDF** para estudiantes de 16 años, aplicando estrategias de simplificación de lenguaje, relevancia para adolescentes y técnicas de engagement.
""")

# Función para extraer texto de un PDF por partes, capítulos y secciones
def extraer_estructura_pdf(archivo_pdf):
    reader = PyPDF2.PdfReader(archivo_pdf)
    texto_completo = ""
    for page in reader.pages:
        texto_completo += page.extract_text() + "\n"

    # Definir patrones para partes, capítulos y secciones
    patron_parte = re.compile(r'Parte\s+\w+', re.IGNORECASE)
    patron_capitulo = re.compile(r'Capítulo\s+\d+', re.IGNORECASE)
    patron_seccion = re.compile(r'Sección\s+\d+', re.IGNORECASE)

    # Encontrar todas las partes
    partes = patron_parte.findall(texto_completo)
    if partes:
        # Estructura jerárquica con partes, capítulos y secciones
        estructura = {}
        indices_partes = [m.start() for m in patron_parte.finditer(texto_completo)]
        indices_partes.append(len(texto_completo))  # Para capturar el último segmento

        for i in range(len(partes)):
            inicio = indices_partes[i]
            fin = indices_partes[i+1]
            texto_parte = texto_completo[inicio:fin]
            titulo_parte = partes[i].strip()
            estructura[titulo_parte] = {}

            # Encontrar capítulos dentro de la parte
            capitulos = patron_capitulo.findall(texto_parte)
            indices_capitulos = [m.start() for m in patron_capitulo.finditer(texto_parte)]
            indices_capitulos.append(len(texto_parte))  # Para capturar el último segmento

            for j in range(len(capitulos)):
                inicio_cap = indices_capitulos[j]
                fin_cap = indices_capitulos[j+1]
                texto_capitulo = texto_parte[inicio_cap:fin_cap]
                titulo_capitulo = capitulos[j].strip()
                estructura[titulo_parte][titulo_capitulo] = {}

                # Encontrar secciones dentro del capítulo
                secciones = patron_seccion.findall(texto_capitulo)
                indices_secciones = [m.start() for m in patron_seccion.finditer(texto_capitulo)]
                indices_secciones.append(len(texto_capitulo))  # Para capturar el último segmento

                if secciones:
                    for k in range(len(secciones)):
                        inicio_sec = indices_secciones[k]
                        fin_sec = indices_secciones[k+1]
                        texto_seccion = texto_capitulo[inicio_sec:fin_sec]
                        titulo_seccion = secciones[k].strip()
                        estructura[titulo_parte][titulo_capitulo][titulo_seccion] = texto_seccion.replace(titulo_seccion, '').strip()
                else:
                    # Si no hay secciones, asignar el texto completo al capítulo
                    contenido_capitulo = texto_capitulo.replace(titulo_capitulo, '').strip()
                    estructura[titulo_parte][titulo_capitulo] = contenido_capitulo

        return estructura
    else:
        # Si no hay partes, manejar solo capítulos y secciones
        estructura = {}
        capitulos = patron_capitulo.findall(texto_completo)
        indices_capitulos = [m.start() for m in patron_capitulo.finditer(texto_completo)]
        indices_capitulos.append(len(texto_completo))  # Para capturar el último segmento

        for i in range(len(capitulos)):
            inicio = indices_capitulos[i]
            fin = indices_capitulos[i+1]
            texto_capitulo = texto_completo[inicio:fin]
            titulo_capitulo = capitulos[i].strip()
            estructura[titulo_capitulo] = {}

            # Encontrar secciones dentro del capítulo
            secciones = patron_seccion.findall(texto_capitulo)
            indices_secciones = [m.start() for m in patron_seccion.finditer(texto_capitulo)]
            indices_secciones.append(len(texto_capitulo))  # Para capturar el último segmento

            if secciones:
                for j in range(len(secciones)):
                    inicio_sec = indices_secciones[j]
                    fin_sec = indices_secciones[j+1]
                    texto_seccion = texto_capitulo[inicio_sec:fin_sec]
                    titulo_seccion = secciones[j].strip()
                    estructura[titulo_capitulo][titulo_seccion] = texto_seccion.replace(titulo_seccion, '').strip()
            else:
                # Si no hay secciones, asignar el texto completo al capítulo
                contenido_capitulo = texto_capitulo.replace(titulo_capitulo, '').strip()
                estructura[titulo_capitulo] = contenido_capitulo

        return estructura

# Función para adaptar el contenido usando la API
def adaptar_contenido(contenido_original, titulo):
    api_key = st.secrets["api"]["key"]
    url = "https://api.openai.com/v1/chat/completions"  # URL de la API de OpenAI
    
    prompt = f"""
    Adapta el siguiente contenido filosófico para estudiantes de 16 años. Aplica las siguientes estrategias:
    
    **Simplificación del Lenguaje**
    - Reemplaza términos filosóficos complejos con lenguaje cotidiano.
    - Usa oraciones más cortas.
    - Explica conceptos abstractos con analogías relacionadas con la vida diaria de un adolescente.
    - Elimina jerga académica.
    
    **Relevancia para Adolescentes**
    - Conecta las ideas filosóficas con la exploración de la identidad personal, relaciones sociales, desafíos escolares y de la vida, tecnología y experiencias modernas, crecimiento personal e inteligencia emocional.
    
    **Técnicas de Engagement**
    - Incluye un tono conversacional.
    - Usa ejemplos del mundo real.
    - Añade preguntas reflexivas.
    - Desglosa ideas complejas en partes fáciles de entender.
    - Resalta aplicaciones prácticas para la vida diaria.
    
    **Título:**
    {titulo}
    
    **Contenido Original:**
    {contenido_original}
    
    **Adaptación:**
    """

    payload = {
        "messages": [
            {"role": "system", "content": "Eres un asistente que adapta textos filosóficos para estudiantes de 16 años aplicando estrategias de simplificación, relevancia y engagement."},
            {"role": "user", "content": prompt}
        ],
        "model": "gpt-4",
        "temperature": 0.7
    }

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}"
    }

    try:
        response = requests.post(url, headers=headers, json=payload)
        if response.status_code == 200:
            respuesta = response.json()
            adaptacion = respuesta.get('choices')[0].get('message').get('content').strip()
            return adaptacion
        else:
            st.error(f"Error al adaptar {titulo}: {response.status_code} - {response.text}")
            return None
    except Exception as e:
        st.error(f"Excepción al adaptar {titulo}: {e}")
        return None

# Función para convertir texto con formato básico a Python-docx
def texto_a_docx(paragraph, texto):
    # Puedes implementar más reglas de formato si es necesario
    lines = texto.split('\n')
    for line in lines:
        if line.startswith("**") and line.endswith("**"):
            # Texto en negrita
            run = paragraph.add_run(line.strip('*'))
            run.bold = True
        elif line.endswith("?"):
            # Pregunta reflexiva en itálica
            run = paragraph.add_run(line)
            run.italic = True
        else:
            paragraph.add_run(line)
    paragraph.add_run('\n')

# Cargar el archivo PDF
uploaded_file = st.file_uploader("Sube la obra filosófica en PDF", type=["pdf"])

if uploaded_file is not None:
    estructura = extraer_estructura_pdf(uploaded_file)
    if not estructura:
        st.error("No se pudieron extraer partes, capítulos o secciones del PDF. Verifica la estructura del documento.")
    else:
        st.sidebar.header("Configuración de Adaptación")
        opcion = st.sidebar.selectbox(
            "Seleccione una opción de adaptación:",
            ("Seleccionar manualmente", "Adaptar toda la obra")
        )
        
        seleccionados = {}

        if opcion == "Seleccionar manualmente":
            # Navegar por la estructura jerárquica para seleccionar
            for parte, contenido_parte in estructura.items():
                with st.sidebar.expander(parte):
                    for capitulo, contenido_capitulo in contenido_parte.items():
                        if isinstance(contenido_capitulo, dict):
                            for seccion, contenido_seccion in contenido_capitulo.items():
                                key = f"{parte} > {capitulo} > {seccion}"
                                seleccionados[key] = contenido_seccion
                        else:
                            key = f"{parte} > {capitulo}"
                            seleccionados[key] = contenido_capitulo
        else:
            # Adaptar toda la obra
            for parte, contenido_parte in estructura.items():
                if isinstance(contenido_parte, dict):
                    for capitulo, contenido_capitulo in contenido_parte.items():
                        if isinstance(contenido_capitulo, dict):
                            for seccion, contenido_seccion in contenido_capitulo.items():
                                key = f"{parte} > {capitulo} > {seccion}"
                                seleccionados[key] = contenido_seccion
                        else:
                            key = f"{parte} > {capitulo}"
                            seleccionados[key] = contenido_capitulo
                else:
                    seleccionados[parte] = contenido_parte
            st.sidebar.info(f"Se procesarán todas las secciones disponibles: {len(seleccionados)} elementos.")

        # Botón para iniciar la adaptación
        if st.sidebar.button("Adaptar Contenidos"):
            if not seleccionados:
                st.sidebar.error("Por favor, seleccione al menos una parte, capítulo o sección válido.")
            else:
                documentos = {}
                total = len(seleccionados)
                progress_bar = st.progress(0)
                status_text = st.empty()

                for idx, (titulo, contenido) in enumerate(seleccionados.items(), start=1):
                    status_text.text(f"Adaptando {titulo} ({idx}/{total})...")

                    if contenido:
                        adaptacion = adaptar_contenido(contenido, titulo)
                        if adaptacion:
                            documentos[titulo] = adaptacion
                        else:
                            documentos[titulo] = "Error en la adaptación."
                    else:
                        documentos[titulo] = "Contenido no disponible."

                    # Actualizar la barra de progreso
                    progress = idx / total
                    progress_bar.progress(progress)

                    # Simular tiempo de espera (puedes eliminar esto en producción)
                    time.sleep(0.5)

                status_text.text("Adaptación completa.")
                progress_bar.empty()

                # Crear el documento Word
                doc = Document()
                estilo_normal = doc.styles['Normal']
                estilo_normal.font.name = 'Arial'
                estilo_normal.font.size = Pt(12)

                for titulo, contenido in documentos.items():
                    # Agregar título como encabezado de nivel 1
                    doc.add_heading(titulo, level=1)

                    # Dividir el contenido en párrafos basados en saltos de línea
                    paragraphs = contenido.split('\n\n')
                    for para in paragraphs:
                        paragraph = doc.add_paragraph()
                        texto_a_docx(paragraph, para)
                    doc.add_page_break()

                # Guardar el documento en memoria
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                # Botón para descargar
                st.success("Todos los contenidos han sido adaptados exitosamente.")
                st.download_button(
                    label="Descargar Documento Word",
                    data=buffer,
                    file_name="Adapted_Philosophical_Work.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
