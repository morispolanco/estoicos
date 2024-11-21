import streamlit as st
import requests
from bs4 import BeautifulSoup
from docx import Document
from io import BytesIO
import time

# 1. Configuración de la página (Debe ser la primera llamada de Streamlit)
st.set_page_config(
    page_title="Seneca Letters Reimagined",
    layout="wide",
    initial_sidebar_state="expanded",
)

# 2. Inyectar estilos CSS personalizados para una UI más elegante
st.markdown(
    """
    <link href="https://fonts.googleapis.com/css2?family=Roboto:wght@400;700&family=Playfair+Display:wght@700&display=swap" rel="stylesheet">
    <style>
    /* Tipografía global */
    body {
        font-family: 'Roboto', sans-serif;
        background-color: #f0f2f6;
        color: #333333;
    }

    /* Títulos principales */
    .css-1aumxhk h1 {
        color: #2c3e50;
        font-family: 'Playfair Display', serif;
        font-size: 3em;
        text-align: center;
        margin-bottom: 0.5em;
    }

    /* Descripción */
    .css-1d391kg p, .css-1d391kg ul, .css-1d391kg ol {
        font-size: 1.1em;
        line-height: 1.6;
    }

    /* Encabezados secundarios */
    .css-1aumxhk h2, .css-1aumxhk h3 {
        color: #34495e;
        font-family: 'Roboto', sans-serif;
    }

    /* Botones principales */
    .stButton > button {
        background-color: #2980b9;
        color: white;
        border-radius: 8px;
        padding: 12px 24px;
        font-size: 1em;
        transition: background-color 0.3s, transform 0.3s;
    }

    .stButton > button:hover {
        background-color: #1f6391;
        transform: translateY(-2px);
    }

    /* Botón de descarga */
    .stDownloadButton > button {
        background-color: #27ae60;
        color: white;
        border-radius: 8px;
        padding: 12px 24px;
        font-size: 1em;
        transition: background-color 0.3s, transform 0.3s;
    }

    .stDownloadButton > button:hover {
        background-color: #1e8449;
        transform: translateY(-2px);
    }

    /* Barra de progreso */
    div.stProgress > div > div > div {
        background-color: #2980b9;
        border-radius: 10px;
    }

    /* Separadores */
    hr {
        border: 0;
        height: 1px;
        background: #bdc3c7;
        margin: 2em 0;
    }

    /* Footer */
    .footer {
        text-align: center;
        font-size: 0.9em;
        color: #7f8c8d;
        margin-top: 2em;
        padding: 10px 0;
        border-top: 1px solid #bdc3c7;
    }

    /* Formularios en la barra lateral */
    .sidebar .sidebar-content {
        background-color: #ffffff;
        padding: 20px;
        border-radius: 10px;
        box-shadow: 0 4px 12px rgba(0, 0, 0, 0.1);
    }

    /* Mejorar la apariencia de los selects y inputs */
    select, input[type="number"] {
        border: 1px solid #bdc3c7;
        border-radius: 8px;
        padding: 10px;
        width: 100%;
        box-sizing: border-box;
        margin-bottom: 1em;
        font-size: 1em;
        transition: border-color 0.3s;
    }

    select:focus, input[type="number"]:focus {
        border-color: #2980b9;
        outline: none;
        box-shadow: 0 0 5px rgba(41, 128, 185, 0.5);
    }

    /* Tarjetas para cartas adaptadas */
    .card {
        background-color: #ffffff;
        border-radius: 12px;
        box-shadow: 0 4px 12px rgba(0, 0, 0, 0.1);
        padding: 20px;
        margin-bottom: 20px;
        transition: transform 0.2s, box-shadow 0.2s;
    }

    .card:hover {
        transform: translateY(-5px);
        box-shadow: 0 8px 20px rgba(0, 0, 0, 0.15);
    }

    /* Estilo para los mensajes de alerta */
    .stAlert > div {
        border-radius: 8px;
        padding: 15px;
        font-size: 1em;
    }

    .stAlert .stAlert__icon {
        color: #e74c3c; /* Rojo para errores */
    }

    .stSuccess .stAlert__icon {
        color: #27ae60; /* Verde para éxito */
    }

    .stWarning .stAlert__icon {
        color: #f39c12; /* Naranja para advertencias */
    }

    /* Personalizar la barra lateral */
    .sidebar .sidebar-content {
        background-color: #ffffff;
        border-radius: 12px;
        padding: 25px;
        box-shadow: 0 4px 12px rgba(0, 0, 0, 0.1);
    }
    </style>
    """,
    unsafe_allow_html=True
)

# 3. Título de la app
st.title("📜 Reimagine Seneca's Letters for the Modern Corporate World")

# 4. Descripción
st.markdown("""
Bienvenido a la aplicación **Seneca Letters Reimagined**! Transforma la sabiduría atemporal de Seneca en ideas prácticas adaptadas para los gerentes corporativos de hoy en día. Ya sea que estés navegando desafíos de liderazgo, buscando un equilibrio entre el trabajo y la vida personal, o intentando aumentar la productividad, deja que Seneca te guíe a través de las complejidades del entorno laboral moderno.

**Cómo Usar:**
1. Selecciona los números de las cartas que deseas adaptar.
2. Haz clic en "Generar Adaptaciones" para procesar las cartas seleccionadas.
3. Una vez completado, descarga todas tus cartas adaptadas en un único documento Word.
""")

# 5. Inicializar el estado de sesión para almacenar cartas adaptadas
if 'adapted_letters' not in st.session_state:
    st.session_state['adapted_letters'] = {}

# Función para obtener el contenido original de la carta desde Wikisource
def fetch_letter_content(letter_num):
    url = f"https://en.wikisource.org/wiki/Moral_letters_to_Lucilius/Letter_{letter_num}"
    response = requests.get(url)
    
    if response.status_code != 200:
        return None, f"Carta {letter_num} no encontrada. Asegúrate de que el número de la carta esté entre 1 y 65."
    
    soup = BeautifulSoup(response.content, 'html.parser')
    
    # Extraer el contenido principal de la carta
    content_div = soup.find('div', {'class': 'mw-parser-output'})
    if not content_div:
        return None, "No se pudo analizar el contenido de la carta."
    
    paragraphs = content_div.find_all(['p', 'h2', 'h3'])
    letter_content = ""
    for para in paragraphs:
        if para.name in ['h2', 'h3']:
            continue  # Omitir encabezados
        letter_content += para.get_text(separator="\n") + "\n\n"
    
    if not letter_content.strip():
        return None, "El contenido de la carta está vacío."
    
    return letter_content.strip(), None

# Función para adaptar la carta usando la API de Tune Studio
def adapt_letter(original_content):
    api_url = "https://proxy.tune.app/chat/completions"
    api_key = st.secrets["tune_api_key"]
    
    prompt = f"""
Reimagine the following letter from Seneca to Lucilius, adapting it from its original philosophical content into a modern corporate setting. Imagine Seneca is writing in 2024 and Lucilius is a corporate manager. 

# Original Letter Content:
{original_content}

# Adaptation Instructions:
- Maintain the essence of Seneca's wisdom.
- Use simple, contemporary language.
- Incorporate references to modern technology, work-life balance, stress, leadership, and productivity.
- Adjust metaphors and examples to fit today's corporate culture.
"""
    
    payload = {
        "temperature": 0.66,
        "messages": [
            {
                "role": "user",
                "content": prompt
            }
        ],
        "model": "openai/gpt-4o-mini",
        "stream": False,
        "frequency_penalty": 0,
        "max_tokens": 5194
    }

    headers = {
        "Authorization": api_key,
        "Content-Type": "application/json"
    }

    try:
        response = requests.post(api_url, json=payload, headers=headers)
        response.raise_for_status()
        data = response.json()
        adaptation = data.get('choices', [{}])[0].get('message', {}).get('content', '').strip()
        if not adaptation:
            return None, "No se recibió adaptación de la API."
        return adaptation, None
    except requests.exceptions.RequestException as e:
        return None, f"Fallo en la solicitud a la API: {e}"
    except ValueError:
        return None, "Respuesta inválida de la API."

# Barra lateral para la entrada del usuario
st.sidebar.header("Generar Cartas Adaptadas en Lote")
with st.sidebar.form(key='batch_form'):
    # Permitir selección múltiple de cartas
    letter_numbers = st.multiselect(
        "Selecciona los Números de las Cartas (1-65)",
        options=list(range(1, 66)),
        default=[1]
    )
    submit_button = st.form_submit_button(label='Generar Adaptaciones')

if submit_button:
    if not letter_numbers:
        st.error("Por favor, selecciona al menos una carta para adaptar.")
    else:
        with st.spinner('Obteniendo y adaptando las cartas seleccionadas...'):
            progress_bar = st.progress(0)
            total = len(letter_numbers)
            success_count = 0
            failed_letters = []
            for idx, letter_num in enumerate(letter_numbers):
                original, error = fetch_letter_content(letter_num)
                if error:
                    failed_letters.append((letter_num, error))
                    progress_bar.progress((idx + 1) / total)
                    continue
                adaptation, api_error = adapt_letter(original)
                if api_error:
                    failed_letters.append((letter_num, api_error))
                else:
                    st.session_state['adapted_letters'][letter_num] = adaptation
                    success_count += 1
                progress_bar.progress((idx + 1) / total)
                time.sleep(0.5)  # Pausa para evitar sobrecargar la API

            progress_bar.empty()
            st.success(f"**¡Adaptaciones completadas!** {success_count} de {total} cartas fueron adaptadas exitosamente.")

            if failed_letters:
                st.warning("Algunas cartas no se pudieron adaptar:")
                for num, msg in failed_letters:
                    st.write(f"- **Carta {num}:** {msg}")

# Mostrar las cartas adaptadas
if st.session_state['adapted_letters']:
    st.markdown("---")
    st.header("📝 Tus Cartas Adaptadas")
    sorted_letters = dict(sorted(st.session_state['adapted_letters'].items()))
    for num, content in sorted_letters.items():
        st.markdown(f"<div class='card'><h2>Carta {num}</h2><p>{content.replace('\n', '<br>')}</p></div>", unsafe_allow_html=True)

    # Botón para generar documento Word
    if st.button("📥 Descargar Todas las Cartas Adaptadas como Documento Word"):
        with st.spinner('Generando documento Word...'):
            doc = Document()
            doc.add_heading('Cartas Adaptadas de Seneca a Lucilius', 0)

            for num, content in sorted_letters.items():
                doc.add_heading(f'Carta {num}', level=1)
                doc.add_paragraph(content)
                doc.add_page_break()

            # Guardar el documento en un objeto BytesIO
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            # Proveer botón de descarga
            st.download_button(
                label="✅ Descargar Documento Word",
                data=buffer,
                file_name="Cartas_Adaptadas_Seneca_Letters.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# Pie de página
st.markdown("<div class='footer'>© 2024 Seneca Letters Reimagined | Powered by Streamlit y Tune Studio API</div>", unsafe_allow_html=True)
