import streamlit as st
import requests
from docx import Document
from docx.shared import Pt
from io import BytesIO
import time
from bs4 import BeautifulSoup
import re

# Configuración de la página
st.set_page_config(
    page_title="Adaptador de Cartas de Séneca",
    layout="wide",
    initial_sidebar_state="expanded",
)

# Título de la aplicación
st.title("Adaptador de Cartas de Séneca a un Contexto Corporativo Moderno")

# Descripción
st.markdown("""
Esta aplicación adapta las **65 Cartas de Séneca a Lucilio** a un entorno corporativo moderno, proporcionando consejos relevantes para los gestores de empresas en 2024.
""")

# Función para obtener el contenido de una carta específica
def obtener_contenido_carta(numero):
    url = f"https://en.wikisource.org/wiki/Moral_letters_to_Lucilius/Letter_{numero}"
    response = requests.get(url)
    if response.status_code == 200:
        # Extraer el contenido principal de la carta
        soup = BeautifulSoup(response.content, 'html.parser')
        content_div = soup.find('div', {'class': 'mw-parser-output'})
        if not content_div:
            return None
        paragraphs = content_div.find_all('p')
        contenido = "\n\n".join([para.get_text() for para in paragraphs])
        return contenido
    else:
        return None

# Función para adaptar la carta usando la API de X
def adaptar_carta(contenido_original, numero_carta):
    api_key = st.secrets["api"]["key"]
    url = "https://api.x.ai/v1/chat/completions"  # URL de la API de X

    prompt = f"""
Reimagine letter number {numero_carta} of Seneca to Lucilius, adapting it from its original philosophical content to a modern corporate environment of 2024. Use simple and contemporary language while preserving Seneca's wisdom. Ensure that examples and metaphors are relevant to current challenges faced by corporate managers, including references to modern technology, work-life balance, and common issues such as stress, leadership, and productivity.

Original Content:
{contenido_original}

Adaptation (use proper formatting without Markdown symbols, e.g., use italics instead of asterisks, appropriate heading levels without hashtags):
"""

    payload = {
        "messages": [
            {"role": "system", "content": "You are an assistant that adapts philosophical texts to modern corporate contexts using proper formatting."},
            {"role": "user", "content": prompt}
        ],
        "model": "grok-beta",  # Modelo especificado
        "stream": False,
        "temperature": 0.7
    }

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}"
    }

    try:
        response = requests.post(url, headers=headers, json=payload)
        if response.status_code == 200:
            respuesta = response.json()
            # Asumiendo que la respuesta tiene una estructura similar a OpenAI
            adaptacion = respuesta.get('choices')[0].get('message').get('content').strip()
            return adaptacion
        else:
            st.error(f"Error al adaptar la carta {numero_carta}: {response.status_code} - {response.text}")
            return None
    except Exception as e:
        st.error(f"Excepción al adaptar la carta {numero_carta}: {e}")
        return None

# Función para convertir HTML a formato de Python-docx
def html_a_docx(paragraph, html_text):
    # Parsear el HTML usando BeautifulSoup
    soup = BeautifulSoup(html_text, 'html.parser')
    for elem in soup:
        if isinstance(elem, str):
            paragraph.add_run(elem)
        elif elem.name in ['i', 'em']:
            run = paragraph.add_run(elem.get_text())
            run.italic = True
        elif elem.name in ['b', 'strong']:
            run = paragraph.add_run(elem.get_text())
            run.bold = True
        elif re.match(r'h[1-6]', elem.name):
            # Manejar encabezados
            nivel = int(elem.name[1])
            if 1 <= nivel <= 6:
                paragraph.style = f'Heading {nivel}'
                paragraph.add_run(elem.get_text())
        else:
            # Manejar otros casos o etiquetas desconocidas
            paragraph.add_run(elem.get_text())

# Total de cartas disponibles
TOTAL_CARTAS = 65  # Actualizado a 65 cartas

# Entrada de números de cartas
st.sidebar.header("Configuración de Adaptación")
opcion = st.sidebar.selectbox(
    "Seleccione una opción de cartas a adaptar:",
    ("Seleccionar manualmente", "Adaptar todas las cartas")
)

if opcion == "Seleccionar manualmente":
    cartas_input = st.sidebar.text_input(
        "Ingrese los números de las cartas separados por comas (ejemplo: 1,2,3)",
        value="1"
    )
    # Procesar la entrada
    try:
        numeros_cartas = [int(num.strip()) for num in cartas_input.split(",") if num.strip().isdigit()]
    except:
        numeros_cartas = []
else:
    # Adaptar todas las cartas
    numeros_cartas = list(range(1, TOTAL_CARTAS + 1))
    st.sidebar.info(f"Se procesarán todas las {TOTAL_CARTAS} cartas disponibles.")

# Botón para iniciar la adaptación
if st.sidebar.button("Adaptar Cartas"):
    if not numeros_cartas:
        st.sidebar.error("Por favor, ingrese al menos un número de carta válido.")
    else:
        documentos = {}
        total = len(numeros_cartas)
        progress_bar = st.progress(0)
        status_text = st.empty()

        for idx, numero in enumerate(numeros_cartas, start=1):
            status_text.text(f"Adaptando carta {numero} ({idx}/{total})...")

            contenido_original = obtener_contenido_carta(numero)
            if contenido_original:
                adaptacion = adaptar_carta(contenido_original, numero)
                if adaptacion:
                    documentos[f"Letter {numero}"] = adaptacion
                else:
                    documentos[f"Letter {numero}"] = "Error en la adaptación."
            else:
                documentos[f"Letter {numero}"] = "Contenido no disponible."

            # Actualizar la barra de progreso
            progress = idx / total
            progress_bar.progress(progress)

            # Simular tiempo de espera (puedes eliminar esto en producción)
            time.sleep(0.5)

        status_text.text("Adaptación completa.")
        progress_bar.empty()

        # Crear el documento Word
        doc = Document()
        estilo_normal = doc.styles['Normal']
        estilo_normal.font.name = 'Arial'
        estilo_normal.font.size = Pt(12)

        for titulo, contenido in documentos.items():
            # Agregar título como encabezado de nivel 1
            doc.add_heading(titulo, level=1)

            # Dividir el contenido en párrafos basados en saltos de línea dobles
            paragraphs = contenido.split('\n\n')
            for para in paragraphs:
                paragraph = doc.add_paragraph()
                html_a_docx(paragraph, para)
            doc.add_page_break()

        # Guardar el documento en memoria
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Botón para descargar
        st.success("Todas las cartas han sido adaptadas exitosamente.")
        st.download_button(
            label="Descargar Documento Word",
            data=buffer,
            file_name="Adapted_Seneca_Letters.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
